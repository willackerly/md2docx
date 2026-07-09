#!/usr/bin/env python3
# SPDX-License-Identifier: MIT
"""md2docx — Markdown → styled DOCX, themed by a JSON template.

CONTRACT:C1-THEME-SCHEMA.1.0
CONTRACT:C2-PROVENANCE.1.0

Usage:
  md2docx.py file.md [more.md ...]            # writes <name>.docx next to each source
  md2docx.py -t theme.json file.md            # explicit template
  md2docx.py -o outdir file.md [more.md ...]  # write outputs into a directory
  md2docx.py -o out.docx file.md              # single input, explicit output path
  md2docx.py --no-footer file.md              # suppress footer text

Template resolution (see architecture/CONTRACT-C1-THEME-SCHEMA.1.0.md): the
--template flag wins; otherwise `md2docx-template.json` next to this script;
otherwise `themes/neutral.json` next to this script; otherwise the built-in
neutral defaults below. Any template key may be omitted — it deep-merges
over the defaults.

Markdown support: h1–h4; bullet lists; literal numbered lists; N-column pipe
tables (first row = shaded header); fenced code blocks; inline `code`, **bold**,
*italic*; links (relative links keep their label, absolute URLs appended in
parens); blockquotes; horizontal rules and HTML comments skipped; soft-wrapped
lines joined into one paragraph. A leading `**CUI...**`-style banner line is
promoted to the page header and footer (marking convention) and replaces the
footer text — see the theme's `cui_banner` block and `themes/marked-docs.json`
for an example.

Requires: python-docx  (pip install python-docx)
"""
import argparse
import datetime
import hashlib
import json
import re
import sys
from pathlib import Path

TOOL_VERSION = "0.2.0"

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEFAULTS = {
    "fonts": {"body": "Calibri", "mono": "Consolas"},
    "base": {"size_pt": 10.5, "color": "1A1A1A", "space_after_pt": 6},
    "headings": {
        "h1": {"size_pt": 18, "color": "000000"},
        "h2": {"size_pt": 14, "color": "000000"},
        "h3": {"size_pt": 11.5, "color": "000000"},
        "h4": {"size_pt": 10.5, "color": "333333"},
        "bold": True,
        "space_before_pt": 12,
        "space_after_pt": 5,
    },
    "table": {
        "style": "Table Grid",
        "header_fill": "EEEEEE",
        "header_color": "222222",
        "header_bold": True,
        "cell_size_pt": 9.5,
    },
    "code": {"fill": "F5F5F5", "color": "333333", "block_size_pt": 8.5},
    "blockquote": {
        "color": "555555",
        "size_pt": 9.5,
        "border_color": "888888",
        "border_size_eighths": 18,
        "indent_in": 0.3,
    },
    "footer": {"text": "", "color": "999999", "size_pt": 10, "bold": True},
    "cui_banner": {"detect": True, "size_pt": 10, "bold": True, "color": "000000"},
}


def deep_merge(base, over):
    out = dict(base)
    for k, v in (over or {}).items():
        out[k] = deep_merge(base[k], v) if isinstance(v, dict) and isinstance(base.get(k), dict) else v
    return out


def rgb(hexstr):
    return RGBColor(*(int(hexstr[i:i + 2], 16) for i in (0, 2, 4)))


def _shd(el_get, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    el_get().append(shd)


def shade_cell(cell, fill): _shd(cell._tc.get_or_add_tcPr, fill)
def shade_para(p, fill): _shd(p._p.get_or_add_pPr, fill)
def shade_run(r, fill): _shd(r._r.get_or_add_rPr, fill)


def left_border(p, color, size_eighths):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_eighths))
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def demote_links(text):
    def repl(m):
        label, url = m.group(1), m.group(2)
        if url.startswith("http") and label != url:
            return f"{label} ({url})"
        return label
    return LINK_RE.sub(repl, text)


def stamp_provenance(doc, src, tpl_path, cfg):
    """Embed round-trip breadcrumbs in DOCX core properties.

    CONTRACT:C2-PROVENANCE.1.0 — see architecture/CONTRACT-C2-PROVENANCE.1.0.md
    for the field table and the 255-char / hash-truncation rules this
    function must honor.

    Survives Word saves (core props are standard OPC parts). docx2md reads
    `comments` to recover which template produced the formatting and whether
    the source MD has drifted since conversion (hash comparison).
    """
    src_bytes = Path(src).read_bytes()
    tpl_sha = (hashlib.sha256(Path(tpl_path).read_bytes()).hexdigest()[:16]
               if tpl_path and Path(tpl_path).exists() else None)
    prov = {
        "t": f"md2docx/{TOOL_VERSION}",
        "tpl": cfg.get("name") or "built-in defaults",
        "tplsha": tpl_sha,
        "srcsha": hashlib.sha256(src_bytes).hexdigest()[:16],
        "gen": datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%MZ"),
    }
    cp = doc.core_properties                      # each field capped at 255 chars
    cp.comments = json.dumps(prov, separators=(",", ":"))
    cp.subject = str(Path(src).resolve())[-255:]  # full source path, own field
    cp.keywords = "md2docx"
    cp.category = prov["tpl"]


class Converter:
    def __init__(self, cfg, tpl_path=None):
        self.cfg = cfg
        self.tpl_path = tpl_path

    def add_runs(self, p, text, color=None, size=None, base_bold=False):
        cfg = self.cfg
        text = demote_links(text)
        for tok in re.split(r"(`[^`]+`|\*\*.*?\*\*|\*[^*]+\*)", text):
            if not tok:
                continue
            if tok.startswith("`") and tok.endswith("`"):
                r = p.add_run(tok[1:-1])
                r.font.name = cfg["fonts"]["mono"]
                r.font.size = Pt((size or Pt(cfg["base"]["size_pt"])).pt - 0.5)
                r.font.color.rgb = rgb(cfg["code"]["color"])
                shade_run(r, cfg["code"]["fill"])
                if base_bold:
                    r.bold = True
                continue
            if tok.startswith("**") and tok.endswith("**"):
                r = p.add_run(tok[2:-2]); r.bold = True
            elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
                r = p.add_run(tok[1:-1]); r.italic = True
            else:
                r = p.add_run(tok)
                if base_bold:
                    r.bold = True
            r.font.name = cfg["fonts"]["body"]
            if color:
                r.font.color.rgb = color
            if size:
                r.font.size = size

    def convert(self, src, out_path, footer_on=True):
        cfg = self.cfg
        lines = Path(src).read_text(encoding="utf8").splitlines()
        doc = Document()

        n = doc.styles["Normal"]
        n.font.name = cfg["fonts"]["body"]
        n.font.size = Pt(cfg["base"]["size_pt"])
        n.font.color.rgb = rgb(cfg["base"]["color"])
        n.paragraph_format.space_after = Pt(cfg["base"]["space_after_pt"])
        for lvl in (1, 2, 3, 4):
            h = cfg["headings"][f"h{lvl}"]
            s = doc.styles[f"Heading {lvl}"]
            s.font.name = cfg["fonts"]["body"]
            s.font.size = Pt(h["size_pt"])
            s.font.color.rgb = rgb(h["color"])
            s.font.bold = cfg["headings"]["bold"]
            s.paragraph_format.space_before = Pt(cfg["headings"]["space_before_pt"])
            s.paragraph_format.space_after = Pt(cfg["headings"]["space_after_pt"])

        banner = None
        body_start = 0
        if cfg["cui_banner"]["detect"]:
            for idx, ln in enumerate(lines):
                if not ln.strip():
                    continue
                m = re.fullmatch(r"\*\*(CUI[^*]*)\*\*", ln.strip())
                if m:
                    banner = m.group(1)
                    body_start = idx + 1
                break

        i = body_start
        while i < len(lines):
            ln = lines[i].rstrip()
            stripped = ln.strip()
            if not stripped or stripped.startswith("<!--") or re.fullmatch(r"-{3,}", stripped):
                i += 1; continue

            if stripped.startswith("```"):
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Inches(0.2)
                    shade_para(p, cfg["code"]["fill"])
                    r = p.add_run(lines[i] if lines[i].strip() else " ")
                    r.font.name = cfg["fonts"]["mono"]
                    r.font.size = Pt(cfg["code"]["block_size_pt"])
                    r.font.color.rgb = rgb(cfg["code"]["color"])
                    i += 1
                i += 1
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue

            if stripped.startswith("|"):
                rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    if not all(re.fullmatch(r":?-+:?", c) for c in cells):
                        rows.append(cells)
                    i += 1
                if rows:
                    tc = cfg["table"]
                    ncols = max(len(r) for r in rows)
                    t = doc.add_table(rows=0, cols=ncols)
                    t.style = tc["style"]
                    t.autofit = True
                    for ri, row in enumerate(rows):
                        cells = t.add_row().cells
                        for ci in range(ncols):
                            txt = row[ci] if ci < len(row) else ""
                            p = cells[ci].paragraphs[0]
                            if ri == 0:
                                self.add_runs(p, re.sub(r"\*\*", "", txt),
                                              color=rgb(tc["header_color"]),
                                              size=Pt(tc["cell_size_pt"]),
                                              base_bold=tc["header_bold"])
                                shade_cell(cells[ci], tc["header_fill"])
                            else:
                                self.add_runs(p, txt, size=Pt(tc["cell_size_pt"]))
                    doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue

            if stripped.startswith(">"):
                buf = []
                while i < len(lines) and lines[i].strip().startswith(">"):
                    buf.append(lines[i].strip().lstrip(">").strip())
                    i += 1
                bq = cfg["blockquote"]
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(bq["indent_in"])
                left_border(p, bq["border_color"], bq["border_size_eighths"])
                self.add_runs(p, " ".join(buf), color=rgb(bq["color"]), size=Pt(bq["size_pt"]))
                continue

            m = re.match(r"^(#{1,4}) ", stripped)
            if m:
                doc.add_heading(stripped[len(m.group(1)) + 1:], level=len(m.group(1)))
            elif re.match(r"^[-*] ", stripped):
                p = doc.add_paragraph(style="List Bullet")
                self.add_runs(p, stripped[2:])
            elif re.match(r"^\d+\. ", stripped):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                self.add_runs(p, stripped)
            else:
                buf = [stripped]
                while i + 1 < len(lines):
                    nxt = lines[i + 1].strip()
                    if (not nxt or nxt.startswith(("#", "|", ">", "```", "<!--"))
                            or re.match(r"^[-*] ", nxt) or re.match(r"^\d+\. ", nxt)
                            or re.fullmatch(r"-{3,}", nxt)):
                        break
                    buf.append(nxt); i += 1
                p = doc.add_paragraph()
                self.add_runs(p, " ".join(buf))
            i += 1

        sec = doc.sections[0]
        if banner:
            cb = cfg["cui_banner"]
            hp = sec.header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr = hp.add_run(banner)
            hr.font.name = cfg["fonts"]["body"]
            hr.font.size = Pt(cb["size_pt"])
            hr.bold = cb["bold"]
            hr.font.color.rgb = rgb(cb["color"])
        ftext = banner if banner else (cfg["footer"]["text"] if footer_on else "")
        if ftext:
            fc = cfg["footer"]
            fp = sec.footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = fp.add_run(ftext)
            fr.font.name = cfg["fonts"]["body"]
            fr.font.size = Pt(fc["size_pt"])
            fr.bold = fc["bold"]
            if not banner:
                fr.font.color.rgb = rgb(fc["color"])

        stamp_provenance(doc, src, self.tpl_path, self.cfg)
        doc.save(out_path)
        return out_path


def resolve_template(explicit, base_dir=None):
    """Template resolution order — CONTRACT:C1-THEME-SCHEMA.1.0.

    1. --template flag, if given.
    2. `md2docx-template.json` next to this script (a local override; not
       shipped in this repo).
    3. `themes/neutral.json` next to this script (the built-in-equivalent
       shipped theme, kept as a JSON file so it's documented and diffable).
    4. The DEFAULTS dict in this module, hard-coded.

    `base_dir` overrides "next to this script" (defaults to this module's
    directory); tests use it to exercise the resolution order against a
    scratch directory instead of the real repo.

    Returns (cfg, tpl_path, message) — tpl_path is None when no file was
    used (pure DEFAULTS), so provenance stamping records "built-in defaults".
    """
    here = Path(base_dir) if base_dir else Path(__file__).parent
    if explicit:
        tpl_path = Path(explicit)
        if not tpl_path.exists():
            sys.exit(f"template not found: {explicit}")
        cfg = deep_merge(DEFAULTS, json.loads(tpl_path.read_text(encoding="utf8")))
        return cfg, tpl_path, f"template: {tpl_path}" + (f" ({cfg.get('name')})" if cfg.get("name") else "")

    local_override = here / "md2docx-template.json"
    if local_override.exists():
        cfg = deep_merge(DEFAULTS, json.loads(local_override.read_text(encoding="utf8")))
        return cfg, local_override, f"template: {local_override}" + (f" ({cfg.get('name')})" if cfg.get("name") else "")

    shipped_neutral = here / "themes" / "neutral.json"
    if shipped_neutral.exists():
        cfg = deep_merge(DEFAULTS, json.loads(shipped_neutral.read_text(encoding="utf8")))
        return cfg, shipped_neutral, f"template: {shipped_neutral} ({cfg.get('name', 'Neutral')})"

    return DEFAULTS, None, "template: built-in defaults (no md2docx-template.json or themes/neutral.json found)"


def main():
    ap = argparse.ArgumentParser(description="Markdown → styled DOCX, themed by a JSON template.")
    ap.add_argument("files", nargs="+", help="Markdown source files")
    ap.add_argument("-t", "--template", help="Template JSON (default: md2docx-template.json next to this script, then themes/neutral.json)")
    ap.add_argument("-o", "--out", help="Output directory, or output .docx path for a single input")
    ap.add_argument("--no-footer", action="store_true", help="Suppress footer text")
    args = ap.parse_args()

    cfg, tpl_path, message = resolve_template(args.template)
    print(message)

    conv = Converter(cfg, tpl_path=tpl_path)
    single_file_out = args.out and args.out.lower().endswith(".docx")
    if single_file_out and len(args.files) > 1:
        sys.exit("-o <file.docx> only valid with a single input; use -o <dir> for many")

    for src in args.files:
        src_p = Path(src)
        if not src_p.exists():
            sys.exit(f"not found: {src}")
        if single_file_out:
            out = Path(args.out)
        elif args.out:
            Path(args.out).mkdir(parents=True, exist_ok=True)
            out = Path(args.out) / (src_p.stem + ".docx")
        else:
            out = src_p.with_suffix(".docx")
        print("→", conv.convert(src_p, out, footer_on=not args.no_footer))


if __name__ == "__main__":
    main()

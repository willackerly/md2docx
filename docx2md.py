#!/usr/bin/env python3
# SPDX-License-Identifier: MIT
"""docx2md — styled DOCX -> canonical Markdown (reverse of md2docx.py).

CONTRACT:C3-ROUNDTRIP.1.0

Style-driven-inversion converter: it walks the document body in order and
inverts the forward converter's construct->style mapping:

  Heading 1-4 style        -> #, ##, ###, ####
  List Bullet style        -> "- "
  Normal para "N. ..."     -> kept literally (forward writes literal numbers)
  table (row 0 shaded)     -> pipe table with a separator row
  left-bordered indent     -> "> " blockquote
  paragraph-shaded + mono  -> ``` fenced code block
  inline: mono/shaded run  -> `code`,  bold run -> **, italic run -> *

Adjacent runs with identical formatting are merged before emitting so words
are not fragmented. Output is canonical MD: "-" bullets, "**" bold, one blank
line between blocks, no trailing whitespace. See
architecture/CONTRACT-C3-ROUNDTRIP.1.0.md for the full round-trip guarantee
and what forward conversion drops.

Not handled (out of scope for this tier): custom-XML-part source recovery and
3-way merge; recovering link markup the forward tool demoted to plain text.

Usage:
  docx2md.py file.docx [-o out.md]      # default output next to input
Requires: python-docx
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

MONO_FONTS = {"Consolas", "Courier New", "Courier", "Menlo", "Monaco",
              "DejaVu Sans Mono", "Cascadia Code", "Cascadia Mono"}


# ---------- low-level XML probes ----------

def _pPr_child(el, tag):
    pPr = el.find(qn("w:pPr"))
    return None if pPr is None else pPr.find(qn(tag))


def para_shd_fill(p):
    """Paragraph-level shading fill (marks a code-block line), or None."""
    shd = _pPr_child(p._p, "w:shd")
    return shd.get(qn("w:fill")) if shd is not None else None


def para_has_left_border(p):
    """True if the paragraph carries a left border (marks a blockquote)."""
    pBdr = _pPr_child(p._p, "w:pBdr")
    return pBdr is not None and pBdr.find(qn("w:left")) is not None


def run_shd_fill(r):
    rPr = r._r.find(qn("w:rPr"))
    if rPr is None:
        return None
    shd = rPr.find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def run_is_code(r):
    return run_shd_fill(r) is not None or (r.font.name in MONO_FONTS)


# ---------- inline run inversion ----------

def _wrap(text, marker):
    """Wrap non-space core in marker, keeping surrounding whitespace outside."""
    core = text.strip()
    if not core:
        return text
    lead = text[:len(text) - len(text.lstrip())]
    trail = text[len(text.rstrip()):]
    return f"{lead}{marker}{core}{marker}{trail}"


def runs_to_md(runs, suppress_bold=False):
    """Invert a run sequence to inline Markdown, merging like-formatted runs."""
    segs = []  # (fmt, text) where fmt = (code, bold, italic)
    for r in runs:
        if not r.text:
            continue
        code = run_is_code(r)
        bold = bool(r.bold) and not suppress_bold
        italic = bool(r.italic)
        if code:  # code is exclusive; bold/italic markers do not apply inside
            bold = italic = False
        fmt = (code, bold, italic)
        if segs and segs[-1][0] == fmt:
            segs[-1] = (fmt, segs[-1][1] + r.text)
        else:
            segs.append((fmt, r.text))

    out = []
    for (code, bold, italic), text in segs:
        if code:
            out.append("`" + text + "`")
            continue
        if bold and italic:
            out.append(_wrap(text, "***"))
        elif bold:
            out.append(_wrap(text, "**"))
        elif italic:
            out.append(_wrap(text, "*"))
        else:
            out.append(text)
    return "".join(out)


# ---------- cell / table helpers ----------

def cell_md(cell, header=False):
    parts = [runs_to_md(p.runs, suppress_bold=header) for p in cell.paragraphs]
    txt = " ".join(s for s in (p.strip() for p in parts) if s)
    return txt.replace("|", "\\|")


def table_to_md(tbl):
    rows = tbl.rows
    if not rows:
        return ""
    ncols = len(tbl.columns)
    lines = []
    header = ["" for _ in range(ncols)]
    for ci, cell in enumerate(rows[0].cells[:ncols]):
        header[ci] = cell_md(cell, header=True)
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * ncols) + " |")
    for row in rows[1:]:
        vals = ["" for _ in range(ncols)]
        for ci, cell in enumerate(row.cells[:ncols]):
            vals[ci] = cell_md(cell)
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


# ---------- provenance ----------

def print_provenance(doc):
    """Recover the CONTRACT:C2-PROVENANCE.1.0 stamp written by md2docx.py."""
    cp = doc.core_properties
    print("--- provenance stamp ---", file=sys.stderr)
    if cp.comments:
        try:
            prov = json.loads(cp.comments)
            for k, v in prov.items():
                print(f"  {k}: {v}", file=sys.stderr)
        except (ValueError, TypeError):
            print(f"  comments: {cp.comments}", file=sys.stderr)
    else:
        print("  (no comments JSON found)", file=sys.stderr)
    if cp.subject:
        print(f"  subject (source path): {cp.subject}", file=sys.stderr)
    if cp.category:
        print(f"  category (template):   {cp.category}", file=sys.stderr)
    if cp.keywords:
        print(f"  keywords: {cp.keywords}", file=sys.stderr)
    print("------------------------", file=sys.stderr)


# ---------- main body walk ----------

def heading_level(style_name):
    m = re.fullmatch(r"Heading (\d+)", style_name or "")
    return int(m.group(1)) if m else None


def convert(docx_path):
    doc = Document(str(docx_path))
    print_provenance(doc)

    blocks = []

    # Page-header banner (e.g. a CUI-style marking) -> leading **BANNER** line.
    sec = doc.sections[0]
    hdr = sec.header.paragraphs[0].text.strip() if sec.header.paragraphs else ""
    if hdr:
        blocks.append(f"**{hdr}**")

    code_buf = None  # accumulating fenced-code lines

    def flush_code():
        nonlocal code_buf
        if code_buf is not None:
            body = "\n".join(l.rstrip() for l in code_buf)
            blocks.append("```\n" + body + "\n```")
            code_buf = None

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]

        if tag == "tbl":
            flush_code()
            md = table_to_md(Table(child, doc))
            if md:
                blocks.append(md)
            continue

        if tag != "p":
            continue

        p = Paragraph(child, doc)

        # Code-block line: paragraph shading.
        if para_shd_fill(p) is not None:
            line = p.runs[0].text if p.runs else ""
            if code_buf is None:
                code_buf = []
            code_buf.append(line)
            continue

        flush_code()

        style = p.style.name if p.style else "Normal"
        text = runs_to_md(p.runs)

        # Blank spacer paragraphs the forward tool inserts between blocks.
        if not text.strip():
            continue

        lvl = heading_level(style)
        if lvl:
            blocks.append("#" * lvl + " " + text.strip())
        elif style == "List Bullet":
            blocks.append("- " + text.strip())
        elif para_has_left_border(p):
            blocks.append("> " + text.strip())
        else:
            blocks.append(text.rstrip())

    flush_code()

    md = "\n\n".join(b.rstrip() for b in blocks)
    md = "\n".join(line.rstrip() for line in md.split("\n"))
    return md.rstrip() + "\n"


def main():
    ap = argparse.ArgumentParser(description="Styled DOCX -> canonical Markdown.")
    ap.add_argument("file", help="Input .docx")
    ap.add_argument("-o", "--out", help="Output .md (default: next to input)")
    args = ap.parse_args()

    src = Path(args.file)
    if not src.exists():
        sys.exit(f"not found: {src}")
    out = Path(args.out) if args.out else src.with_suffix(".md")

    md = convert(src)
    out.write_text(md, encoding="utf8")
    print(f"-> {out}")


if __name__ == "__main__":
    main()
